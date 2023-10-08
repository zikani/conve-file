import json
import csv
import yaml
import xml.etree.ElementTree as ET
from openpyxl import load_workbook
from docx import Document
import PyPDF2
from ebooklib import epub
import xlrd
from bs4 import BeautifulSoup
from fpdf import FPDF
from latex import build_pdf

class DataConverter:
    @staticmethod
    def json_to_csv(json_file, csv_file):
        # Conversion code for JSON to CSV
        try:
            with open(json_file, 'r') as json_file:
                data = json.load(json_file)

            with open(csv_file, mode='w', newline='') as csv_file:
                writer = csv.DictWriter(csv_file, fieldnames=data[0].keys())
                writer.writeheader()
                
                for row in data:
                    writer.writerow(row)

            return f"Successfully converted JSON to CSV and saved to {csv_file}"
        except FileNotFoundError:
            return f"Input file '{json_file}' not found."
        except (json.JSONDecodeError, KeyError) as e:
            return f"Error converting JSON to CSV: {str(e)}"

    @staticmethod
    def csv_to_json(csv_file, json_file):
        # Conversion code for CSV to JSON
        try:
            data = []

            with open(csv_file, 'r') as csv_file:
                reader = csv.DictReader(csv_file)
                for row in reader:
                    data.append(row)

            with open(json_file, 'w') as json_file:
                json.dump(data, json_file, indent=4)

            return f"Successfully converted CSV to JSON and saved to {json_file}"
        except FileNotFoundError:
            return f"Input file '{csv_file}' not found."
        except Exception as e:
            return f"Error converting CSV to JSON: {str(e)}"

    @staticmethod
    def json_to_yaml(json_file, yaml_file):
        # Conversion code for JSON to YAML
        try:
            with open(json_file, 'r') as json_file:
                data = json.load(json_file)

            with open(yaml_file, 'w') as yaml_file:
                yaml.dump(data, yaml_file, default_flow_style=False)

            return f"Successfully converted JSON to YAML and saved to {yaml_file}"
        except FileNotFoundError:
            return f"Input file '{json_file}' not found."
        except (json.JSONDecodeError, KeyError) as e:
            return f"Error converting JSON to YAML: {str(e)}"

    @staticmethod
    def xml_to_json(xml_file, json_file):
        # Conversion code for XML to JSON
        try:
            tree = ET.parse(xml_file)
            root = tree.getroot()
            data = []

            for element in root:
                entry = {}
                for child in element:
                    entry[child.tag] = child.text
                data.append(entry)

            with open(json_file, 'w') as json_file:
                json.dump(data, json_file, indent=4)

            return f"Successfully converted XML to JSON and saved to {json_file}"
        except FileNotFoundError:
            return f"Input file '{xml_file}' not found."
        except Exception as e:
            return f"Error converting XML to JSON: {str(e)}"

    @staticmethod
    def xlsx_to_json(xlsx_file, json_file):
        # Conversion code for XLSX to JSON
        try:
            data = []

            workbook = load_workbook(xlsx_file)
            sheet = workbook.active

            for row in sheet.iter_rows(values_only=True):
                entry = {}
                for col_num, cell_value in enumerate(row):
                    entry[sheet.cell(1, col_num + 1).value] = cell_value
                data.append(entry)

            with open(json_file, 'w') as json_file:
                json.dump(data, json_file, indent=4)

            return f"Successfully converted XLSX to JSON and saved to {json_file}"
        except FileNotFoundError:
            return f"Input file '{xlsx_file}' not found."
        except Exception as e:
            return f"Error converting XLSX to JSON: {str(e)}"

    @staticmethod
    def docx_to_json(docx_file, json_file):
        # Conversion code for DOCX to JSON
        try:
            data = []

            doc = Document(docx_file)

            for paragraph in doc.paragraphs:
                entry = {}
                for run in paragraph.runs:
                    entry[run.bold] = run.text
                data.append(entry)

            with open(json_file, 'w') as json_file:
                json.dump(data, json_file, indent=4)

            return f"Successfully converted DOCX to JSON and saved to {json_file}"
        except FileNotFoundError:
            return f"Input file '{docx_file}' not found."
        except Exception as e:
            return f"Error converting DOCX to JSON: {str(e)}"

    @staticmethod
    def pdf_to_text(pdf_file, text_file):
        # Conversion code for PDF to Text
        try:
            with open(pdf_file, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfFileReader(pdf_file)
                text = ''
                for page_num in range(pdf_reader.numPages):
                    page = pdf_reader.getPage(page_num)
                    text += page.extractText()

            with open(text_file, 'w') as text_file:
                text_file.write(text)

            return f"Successfully converted PDF to text and saved to {text_file}"
        except FileNotFoundError:
            return f"Input file '{pdf_file}' not found."
        except Exception as e:
            return f"Error converting PDF to text: {str(e)}"

    @staticmethod
    def json_to_xml(json_file, xml_file):
        # Conversion code for JSON to XML
        try:
            with open(json_file, 'r') as json_file:
                data = json.load(json_file)

            root = ET.Element('root')

            for entry in data:
                element = ET.Element('item')
                for key, value in entry.items():
                    sub_element = ET.Element(key)
                    sub_element.text = str(value)
                    element.append(sub_element)
                root.append(element)

            tree = ET.ElementTree(root)
            tree.write(xml_file)

            return f"Successfully converted JSON to XML and saved to {xml_file}"
        except FileNotFoundError:
            return f"Input file '{json_file}' not found."
        except (json.JSONDecodeError, KeyError) as e:
            return f"Error converting JSON to XML: {str(e)}"
        except Exception as e:
            return f"Error converting JSON to XML: {str(e)}"

    @staticmethod
    def csv_to_yaml(csv_file, yaml_file):
        # Conversion code for CSV to YAML
        try:
            data = []

            with open(csv_file, 'r') as csv_file:
                reader = csv.DictReader(csv_file)
                for row in reader:
                    data.append(row)

            with open(yaml_file, 'w') as yaml_file:
                yaml.dump(data, yaml_file, default_flow_style=False)

            return f"Successfully converted CSV to YAML and saved to {yaml_file}"
        except FileNotFoundError:
            return f"Input file '{csv_file}' not found."
        except Exception as e:
            return f"Error converting CSV to YAML: {str(e)}"
        
    @staticmethod
    def docx_to_txt(docx_file, txt_file):
        # Conversion code for DOCX to TXT
        try:
            doc = Document(docx_file)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            with open(txt_file, 'w') as txt_file:
                txt_file.write(text)

            return f"Successfully converted DOCX to TXT and saved to {txt_file}"
        except FileNotFoundError:
            return f"Input file '{docx_file}' not found."
        except Exception as e:
            return f"Error converting DOCX to TXT: {str(e)}"

    @staticmethod
    def html_to_txt(html_file, txt_file):
        # Conversion code for HTML to TXT
        try:
            with open(html_file, 'r', encoding='utf-8') as html_file:
                soup = BeautifulSoup(html_file, 'html.parser')
                text = soup.get_text()

            with open(txt_file, 'w') as txt_file:
                txt_file.write(text)

            return f"Successfully converted HTML to TXT and saved to {txt_file}"
        except FileNotFoundError:
            return f"Input file '{html_file}' not found."
        except Exception as e:
            return f"Error converting HTML to TXT: {str(e)}"

    @staticmethod
    def epub_to_txt(epub_file, txt_file):
        # Conversion code for EPUB to TXT
        try:
            book = epub.read_epub(epub_file)
            text = ""

            for item in book.get_items():
                if item.get_type() == epub.EpubHtml:
                    text += item.get_content() + "\n"

            with open(txt_file, 'w') as txt_file:
                txt_file.write(text)

            return f"Successfully converted EPUB to TXT and saved to {txt_file}"
        except FileNotFoundError:
            return f"Input file '{epub_file}' not found."
        except Exception as e:
            return f"Error converting EPUB to TXT: {str(e)}"

    @staticmethod
    def latex_to_txt(latex_file, txt_file):
        # Conversion code for LaTeX to TXT
        try:
            with open(latex_file, 'r', encoding='utf-8') as latex_file:
                latex_content = latex_file.read()

                # You may need to use a LaTeX-to-Text conversion library here

            with open(txt_file, 'w') as txt_file:
                txt_file.write(latex_content)

            return f"Successfully converted LaTeX to TXT and saved to {txt_file}"
        except FileNotFoundError:
            return f"Input file '{latex_file}' not found."
        except Exception as e:
            return f"Error converting LaTeX to TXT: {str(e)}"

    
    @staticmethod
    def xls_to_txt(xls_file, txt_file):
        # Conversion code for XLS to TXT
        try:
            workbook = xlrd.open_workbook(xls_file)
            sheet = workbook.sheet_by_index(0)
            text = ""

            for row_num in range(sheet.nrows):
                row = sheet.row_values(row_num)
                text += '\t'.join(map(str, row)) + '\n'

            with open(txt_file, 'w') as txt_file:
                txt_file.write(text)

            return f"Successfully converted XLS to TXT and saved to {txt_file}"
        except FileNotFoundError:
            return f"Input file '{xls_file}' not found."
        except Exception as e:
            return f"Error converting XLS to TXT: {str(e)}"

    @staticmethod
    def txt_to_txt(txt_file, txt_output_file):
        # Conversion code for TXT to TXT (no conversion needed)
        try:
            with open(txt_file, 'r', encoding='utf-8') as txt_file:
                txt_content = txt_file.read()

            with open(txt_output_file, 'w') as txt_output_file:
                txt_output_file.write(txt_content)

            return f"Successfully converted TXT to TXT and saved to {txt_output_file}"
        except FileNotFoundError:
            return f"Input file '{txt_file}' not found."
        except Exception as e:
            return f"Error converting TXT to TXT: {str(e)}"

    @staticmethod
    def json_to_json(json_input_file, json_output_file):
        # Conversion code for JSON to JSON (no conversion needed)
        try:
            with open(json_input_file, 'r') as json_file:
                json_content = json_file.read()

            with open(json_output_file, 'w') as json_output_file:
                json_output_file.write(json_content)

            return f"Successfully converted JSON to JSON and saved to {json_output_file}"
        except FileNotFoundError:
            return f"Input file '{json_input_file}' not found."
        except Exception as e:
            return f"Error converting JSON to JSON: {str(e)}"

def main():
    # Get user input for input and output file paths and conversion types
    input_format = input("Enter the input format (JSON, CSV, XML, XLSX, docx, PDF, HTML, ODT, EPUB, LaTeX, TXT, XLS, ODS): ").strip().lower()
    output_format = input("Enter the output format (JSON, CSV, YAML, XLSX, docx, PDF, Text, TXT, JSON): ").strip().lower()
    input_file = input("Enter the path to the input file: ")
    output_file = input("Enter the path to the output file: ")

    converter = DataConverter()

    # Perform the selected conversion
    if input_format == 'json' and output_format == 'csv':
        result = converter.json_to_csv(input_file, output_file)
    elif input_format == 'csv' and output_format == 'json':
        result = converter.csv_to_json(input_file, output_file)
    elif input_format == 'json' and output_format == 'yaml':
        result = converter.json_to_yaml(input_file, output_file)
    elif input_format == 'xml' and output_format == 'json':
        result = converter.xml_to_json(input_file, output_file)
    elif input_format == 'xlsx' and output_format == 'json':
        result = converter.xlsx_to_json(input_file, output_file)
    elif input_format == 'docx' and output_format == 'json':
        result = converter.docx_to_json(input_file, output_file)
    elif input_format == 'pdf' and output_format == 'text':
        result = converter.pdf_to_text(input_file, output_file)
    elif input_format == 'json' and output_format == 'xml':
        result = converter.json_to_xml(input_file, output_file)  
    elif input_format == 'csv' and output_format == 'yaml':
        result = converter.csv_to_yaml(input_file, output_file)  
    elif input_format == 'docx' and output_format == 'txt':
        result = converter.docx_to_txt(input_file, output_file)
    elif input_format == 'html' and output_format == 'txt':
        result = converter.html_to_txt(input_file, output_file)
    elif input_format == 'epub' and output_format == 'txt':
        result = converter.epub_to_txt(input_file, output_file)
    elif input_format == 'latex' and output_format == 'txt':
        result = converter.latex_to_txt(input_file, output_file)
    elif input_format == 'txt' and output_format == 'txt':
        result = converter.txt_to_txt(input_file, output_file)
    elif input_format == 'xls' and output_format == 'txt':
        result = converter.xls_to_txt(input_file, output_file)
    elif input_format == 'json' and output_format == 'json':
        result = converter.json_to_json(input_file, output_file)
    else:
        result = "Unsupported conversion format."

    print(result)

if __name__ == "__main__":
    main()
